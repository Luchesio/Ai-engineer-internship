import csv
import io
import json
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx", ".csv", ".json"}

CHUNK_WORDS = 700
CHUNK_OVERLAP = 80


class UnsupportedDocument(ValueError):
    pass


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _from_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _from_csv(data: bytes) -> str:
    rows = list(csv.reader(io.StringIO(data.decode("utf-8", errors="replace"))))
    return "\n".join(" | ".join(row) for row in rows)


def _from_json(data: bytes) -> str:
    return json.dumps(json.loads(data.decode("utf-8", errors="replace")), indent=2)


def _from_plaintext(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


EXTRACTORS = {
    ".txt": _from_plaintext,
    ".md": _from_plaintext,
    ".markdown": _from_plaintext,
    ".pdf": _from_pdf,
    ".docx": _from_docx,
    ".csv": _from_csv,
    ".json": _from_json,
}


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in EXTRACTORS:
        raise UnsupportedDocument(
            f"'{suffix or filename}' is not supported. Use one of: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    text = EXTRACTORS[suffix](data).strip()
    if not text:
        raise UnsupportedDocument(f"No readable text found in '{filename}'. Scanned PDFs need OCR first.")
    return text


def load_path(path: str | Path) -> tuple[str, str]:
    path = Path(path)
    return path.name, extract_text(path.name, path.read_bytes())


def chunk_text(text: str, size: int = CHUNK_WORDS, overlap: int = CHUNK_OVERLAP) -> list[str]:
    words = text.split()
    if len(words) <= size:
        return [text]

    step = max(size - overlap, 1)
    chunks = []
    for start in range(0, len(words), step):
        window = words[start : start + size]
        if not window:
            break
        chunks.append(" ".join(window))
        if start + size >= len(words):
            break
    return chunks